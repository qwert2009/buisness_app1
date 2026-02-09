"""
PDS-Ultimate File Manager
============================
Универсальный файловый менеджер.

По ТЗ:
- Создание файлов ЛЮБЫХ форматов по запросу: Excel, Word, PDF, CSV, TXT, JSON
- Работа с существующими файлами: открыть → прочитать → изменить
- Вести произвольный учёт (не только заказы)
- Объединение данных из нескольких файлов
- Конвертация между форматами
- Архивариус: переименование по стандарту: 2026_02_07_Заказ_Балаклавы.pdf
"""

from __future__ import annotations

import json
import os
from datetime import datetime
from pathlib import Path
from typing import Optional

from pds_ultimate.config import (
    USER_FILES_DIR,
)
from pds_ultimate.core.llm_engine import llm_engine


class FileManager:
    """
    Универсальный файловый менеджер: создание, редактирование,
    конвертация файлов любых форматов.
    """

    def __init__(self, db_session_factory):
        self._session_factory = db_session_factory
        self._excel_engine = ExcelEngine()
        self._pdf_engine = PDFEngine()

    # ═══════════════════════════════════════════════════════════════════════
    # Создание файлов через DeepSeek
    # ═══════════════════════════════════════════════════════════════════════

    async def create_file_from_text(
        self,
        user_request: str,
        file_format: Optional[str] = None,
    ) -> dict:
        """
        Создать файл по текстовому запросу.
        DeepSeek определяет структуру, менеджер генерирует файл.

        «Создай таблицу учёта сотрудников»
        → DeepSeek → структура → Excel
        """
        # Определить формат
        if not file_format:
            file_format = await self._detect_format(user_request)

        # Получить структуру от DeepSeek
        structure = await self._get_file_structure(user_request, file_format)

        if "error" in structure:
            return structure

        # Генерация файла
        filename = self._generate_filename(
            structure.get("title", "document"), file_format
        )
        filepath = str(USER_FILES_DIR / filename)

        if file_format == "xlsx":
            result = await self._excel_engine.create(filepath, structure)
        elif file_format == "docx":
            result = await self._create_word(filepath, structure)
        elif file_format == "pdf":
            result = await self._pdf_engine.create(filepath, structure)
        elif file_format == "csv":
            result = await self._create_csv(filepath, structure)
        elif file_format == "txt":
            result = await self._create_text(filepath, structure)
        elif file_format == "json":
            result = await self._create_json(filepath, structure)
        else:
            result = await self._create_text(filepath, structure)

        if result.get("success"):
            # Регистрация в БД
            await self._register_file(
                filepath, filename, file_format,
                description=user_request,
            )

        return result

    # ═══════════════════════════════════════════════════════════════════════
    # Работа с существующими файлами
    # ═══════════════════════════════════════════════════════════════════════

    async def read_file(self, filepath: str) -> dict:
        """Прочитать содержимое файла."""
        if not os.path.exists(filepath):
            return {"error": "Файл не найден"}

        ext = Path(filepath).suffix.lower().lstrip(".")

        try:
            if ext in ("xlsx", "xls"):
                return await self._excel_engine.read(filepath)
            elif ext == "docx":
                return await self._read_word(filepath)
            elif ext == "pdf":
                return await self._read_pdf(filepath)
            elif ext == "csv":
                return await self._read_csv(filepath)
            elif ext in ("txt", "md", "json"):
                with open(filepath, "r", encoding="utf-8") as f:
                    content = f.read()
                return {"content": content, "format": ext}
            else:
                return {"error": f"Неподдерживаемый формат: {ext}"}
        except Exception as e:
            return {"error": str(e)}

    async def edit_file(
        self,
        filepath: str,
        instructions: str,
    ) -> dict:
        """
        Редактировать файл по инструкции.
        «Добавь колонку Категория в файл расходов»
        """
        if not os.path.exists(filepath):
            return {"error": "Файл не найден"}

        ext = Path(filepath).suffix.lower().lstrip(".")

        # Прочитать текущее содержимое
        current = await self.read_file(filepath)
        if "error" in current:
            return current

        # Попросить DeepSeek сгенерировать изменения
        prompt = (
            f"Текущее содержимое файла ({ext}):\n"
            f"{json.dumps(current, ensure_ascii=False, default=str)[:3000]}\n\n"
            f"Инструкция: {instructions}\n\n"
            f"Верни JSON с полным обновлённым содержимым файла.\n"
            f'Формат: {{"title":"...","headers":[...],"rows":[...],"content":"..."}}'
        )

        response = await llm_engine.chat(
            message=prompt,
            task_type="analyze",
            temperature=0.2,
            json_mode=True,
        )

        try:
            structure = json.loads(response)
        except Exception:
            return {"error": "DeepSeek не смог сгенерировать изменения"}

        # Перезаписать файл
        if ext == "xlsx":
            return await self._excel_engine.create(filepath, structure)
        elif ext == "docx":
            return await self._create_word(filepath, structure)
        elif ext in ("txt", "md"):
            content = structure.get("content", "")
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(content)
            return {"success": True, "path": filepath}
        elif ext == "json":
            return await self._create_json(filepath, structure)

        return {"error": f"Редактирование {ext} не поддерживается"}

    # ═══════════════════════════════════════════════════════════════════════
    # Конвертация
    # ═══════════════════════════════════════════════════════════════════════

    async def convert_file(
        self,
        source_path: str,
        target_format: str,
    ) -> dict:
        """Конвертировать файл в другой формат."""
        content = await self.read_file(source_path)
        if "error" in content:
            return content

        stem = Path(source_path).stem
        target_name = f"{stem}.{target_format}"
        target_path = str(USER_FILES_DIR / target_name)

        structure = {
            "title": stem,
            "content": content.get("content", ""),
            "headers": content.get("headers", []),
            "rows": content.get("rows", []),
        }

        if target_format == "xlsx":
            return await self._excel_engine.create(target_path, structure)
        elif target_format == "pdf":
            return await self._pdf_engine.create(target_path, structure)
        elif target_format == "csv":
            return await self._create_csv(target_path, structure)
        elif target_format == "txt":
            return await self._create_text(target_path, structure)

        return {"error": f"Конвертация в {target_format} не поддерживается"}

    # ═══════════════════════════════════════════════════════════════════════
    # Архивариус: стандартизация имён
    # ═══════════════════════════════════════════════════════════════════════

    def standardize_filename(
        self,
        original_name: str,
        context: str = "",
    ) -> str:
        """
        Переименовать по стандарту: YYYY_MM_DD_Описание.ext
        «invoice.pdf» → «2026_02_07_Инвойс.pdf»
        """
        today = datetime.now().strftime("%Y_%m_%d")
        ext = Path(original_name).suffix
        stem = Path(original_name).stem

        # Очистка имени
        clean_name = stem.replace(" ", "_").replace("-", "_")
        if context:
            clean_name = context.replace(" ", "_")

        return f"{today}_{clean_name}{ext}"

    # ═══════════════════════════════════════════════════════════════════════
    # Список файлов
    # ═══════════════════════════════════════════════════════════════════════

    async def list_files(
        self,
        category: Optional[str] = None,
    ) -> list[dict]:
        """Список файлов пользователя."""
        from pds_ultimate.core.database import UserFile

        with self._session_factory() as session:
            query = session.query(UserFile)

            if category:
                query = query.filter(UserFile.category == category)

            files = query.order_by(UserFile.created_at.desc()).all()

            return [
                {
                    "id": f.id,
                    "name": f.original_name,
                    "path": f.file_path,
                    "format": f.file_format.value if f.file_format else "?",
                    "size": f.file_size,
                    "category": f.category,
                    "created": f.created_at.isoformat() if f.created_at else None,
                }
                for f in files
            ]

    def format_files_list(self, files: list[dict]) -> str:
        """Форматирование списка файлов."""
        if not files:
            return "📁 Файлов нет."

        lines = [f"📁 Файлы ({len(files)}):\n"]
        for i, f in enumerate(files, 1):
            size_str = ""
            if f.get("size"):
                size_kb = f["size"] / 1024
                size_str = f" ({size_kb:.1f} КБ)"

            lines.append(
                f"  {i}. 📄 {f['name']}{size_str}"
                + (f" [{f['category']}]" if f.get("category") else "")
            )

        return "\n".join(lines)

    # ═══════════════════════════════════════════════════════════════════════
    # Internal: DeepSeek → структура
    # ═══════════════════════════════════════════════════════════════════════

    async def _detect_format(self, request: str) -> str:
        """Определить формат файла из запроса."""
        request_lower = request.lower()

        keywords = {
            "xlsx": ["таблиц", "excel", "xlsx", "учёт", "колонк", "строк"],
            "docx": ["документ", "word", "docx", "письмо", "контракт"],
            "pdf": ["pdf", "инвойс", "счёт", "отчёт"],
            "csv": ["csv"],
            "json": ["json"],
            "txt": ["текст", "заметк", "инструкц"],
        }

        for fmt, kws in keywords.items():
            for kw in kws:
                if kw in request_lower:
                    return fmt

        return "xlsx"  # По умолчанию таблица

    async def _get_file_structure(
        self,
        request: str,
        file_format: str,
    ) -> dict:
        """Получить структуру файла от DeepSeek."""
        prompt = (
            f"Создай структуру файла ({file_format}) по запросу:\n"
            f"«{request}»\n\n"
            f"Верни JSON:\n"
            f'{{"title":"название файла",'
            f'"headers":["Колонка1","Колонка2",...],'
            f'"rows":[["знач1","знач2",...],...],'
            f'"content":"текст если не таблица",'
            f'"sheets":[{{"name":"Лист1","headers":[...],"rows":[...]}}]}}'
        )

        response = await llm_engine.chat(
            message=prompt,
            task_type="analyze",
            temperature=0.3,
            json_mode=True,
        )

        try:
            return json.loads(response)
        except Exception:
            return {"error": "Не удалось определить структуру файла"}

    # ═══════════════════════════════════════════════════════════════════════
    # Internal: Генераторы файлов
    # ═══════════════════════════════════════════════════════════════════════

    async def _create_word(self, filepath: str, structure: dict) -> dict:
        """Создать Word-документ."""
        try:
            from docx import Document

            doc = Document()
            doc.add_heading(structure.get("title", "Документ"), level=1)

            content = structure.get("content", "")
            if content:
                for paragraph in content.split("\n"):
                    if paragraph.strip():
                        doc.add_paragraph(paragraph)

            # Таблица (если есть)
            headers = structure.get("headers", [])
            rows = structure.get("rows", [])

            if headers:
                table = doc.add_table(
                    rows=1 + len(rows), cols=len(headers), style="Table Grid"
                )
                # Заголовки
                for i, h in enumerate(headers):
                    table.rows[0].cells[i].text = str(h)

                # Данные
                for row_idx, row in enumerate(rows, 1):
                    for col_idx, val in enumerate(row):
                        if col_idx < len(headers):
                            table.rows[row_idx].cells[col_idx].text = str(
                                val or "")

            doc.save(filepath)
            return {"success": True, "path": filepath, "format": "docx"}

        except Exception as e:
            return {"error": f"Word creation failed: {e}"}

    async def _create_csv(self, filepath: str, structure: dict) -> dict:
        """Создать CSV-файл."""
        try:
            import csv

            headers = structure.get("headers", [])
            rows = structure.get("rows", [])

            with open(filepath, "w", encoding="utf-8-sig", newline="") as f:
                writer = csv.writer(f)
                if headers:
                    writer.writerow(headers)
                for row in rows:
                    writer.writerow(row)

            return {"success": True, "path": filepath, "format": "csv"}

        except Exception as e:
            return {"error": f"CSV creation failed: {e}"}

    async def _create_text(self, filepath: str, structure: dict) -> dict:
        """Создать текстовый файл."""
        try:
            content = structure.get("content", "")
            if not content:
                # Из таблицы
                lines = []
                if structure.get("title"):
                    lines.append(structure["title"])
                    lines.append("=" * len(structure["title"]))

                headers = structure.get("headers", [])
                rows = structure.get("rows", [])

                if headers:
                    lines.append("\t".join(str(h) for h in headers))
                    lines.append("-" * 40)

                for row in rows:
                    lines.append("\t".join(str(v or "") for v in row))

                content = "\n".join(lines)

            with open(filepath, "w", encoding="utf-8") as f:
                f.write(content)

            return {"success": True, "path": filepath, "format": "txt"}

        except Exception as e:
            return {"error": f"Text creation failed: {e}"}

    async def _create_json(self, filepath: str, structure: dict) -> dict:
        """Создать JSON-файл."""
        try:
            with open(filepath, "w", encoding="utf-8") as f:
                json.dump(structure, f, ensure_ascii=False, indent=2)

            return {"success": True, "path": filepath, "format": "json"}

        except Exception as e:
            return {"error": f"JSON creation failed: {e}"}

    async def _read_word(self, filepath: str) -> dict:
        """Прочитать Word-документ."""
        try:
            from docx import Document

            doc = Document(filepath)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            content = "\n".join(paragraphs)

            return {"content": content, "format": "docx", "paragraphs": len(paragraphs)}

        except Exception as e:
            return {"error": f"Word read failed: {e}"}

    async def _read_pdf(self, filepath: str) -> dict:
        """Прочитать PDF-файл."""
        try:
            import PyPDF2

            with open(filepath, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() or ""

            return {"content": text, "format": "pdf", "pages": len(reader.pages)}

        except Exception as e:
            return {"error": f"PDF read failed: {e}"}

    async def _read_csv(self, filepath: str) -> dict:
        """Прочитать CSV-файл."""
        try:
            import csv

            with open(filepath, "r", encoding="utf-8-sig") as f:
                reader = csv.reader(f)
                rows = list(reader)

            headers = rows[0] if rows else []
            data = rows[1:] if len(rows) > 1 else []

            return {
                "headers": headers,
                "rows": data,
                "format": "csv",
                "row_count": len(data),
            }

        except Exception as e:
            return {"error": f"CSV read failed: {e}"}

    # ═══════════════════════════════════════════════════════════════════════
    # Internal: Регистрация + именование
    # ═══════════════════════════════════════════════════════════════════════

    async def _register_file(
        self,
        filepath: str,
        filename: str,
        file_format: str,
        description: str = "",
        category: Optional[str] = None,
    ) -> None:
        """Зарегистрировать файл в БД."""
        from pds_ultimate.core.database import FileFormat, UserFile

        format_map = {
            "xlsx": FileFormat.XLSX,
            "xls": FileFormat.XLSX,
            "docx": FileFormat.DOCX,
            "pdf": FileFormat.PDF,
            "csv": FileFormat.CSV,
            "txt": FileFormat.TXT,
            "json": FileFormat.JSON,
            "md": FileFormat.MARKDOWN,
        }

        file_size = 0
        if os.path.exists(filepath):
            file_size = os.path.getsize(filepath)

        std_name = self.standardize_filename(filename, description[:50])

        with self._session_factory() as session:
            uf = UserFile(
                original_name=filename,
                standardized_name=std_name,
                file_path=filepath,
                file_format=format_map.get(file_format, FileFormat.TXT),
                file_size=file_size,
                category=category,
                description=description[:500],
            )
            session.add(uf)
            session.commit()

    def _generate_filename(self, title: str, file_format: str) -> str:
        """Сгенерировать имя файла."""
        today = datetime.now().strftime("%Y_%m_%d")
        clean = title.replace(" ", "_").replace("/", "_")[:50]
        return f"{today}_{clean}.{file_format}"


# ═══════════════════════════════════════════════════════════════════════════
# Excel Engine
# ═══════════════════════════════════════════════════════════════════════════

class ExcelEngine:
    """Движок создания/чтения Excel-файлов."""

    async def create(self, filepath: str, structure: dict) -> dict:
        """Создать Excel-файл из структуры."""
        try:
            import xlsxwriter

            wb = xlsxwriter.Workbook(filepath)

            header_fmt = wb.add_format({
                "bold": True,
                "bg_color": "#4472C4",
                "font_color": "#FFFFFF",
                "border": 1,
                "text_wrap": True,
            })
            cell_fmt = wb.add_format({"border": 1, "text_wrap": True})
            money_fmt = wb.add_format({
                "border": 1,
                "num_format": "$#,##0.00",
            })

            sheets = structure.get("sheets", [])
            if not sheets:
                # Одиночный лист
                sheets = [{
                    "name": structure.get("title", "Данные")[:31],
                    "headers": structure.get("headers", []),
                    "rows": structure.get("rows", []),
                }]

            for sheet_data in sheets:
                ws = wb.add_worksheet(sheet_data.get("name", "Лист")[:31])
                headers = sheet_data.get("headers", [])
                rows = sheet_data.get("rows", [])

                # Заголовки
                for col, h in enumerate(headers):
                    ws.write(0, col, str(h), header_fmt)
                    ws.set_column(col, col, max(len(str(h)) + 2, 12))

                # Данные
                for row_idx, row in enumerate(rows, 1):
                    for col_idx, val in enumerate(row):
                        if col_idx < len(headers):
                            # Определить формат
                            try:
                                float_val = float(val)
                                ws.write_number(
                                    row_idx, col_idx, float_val, cell_fmt
                                )
                            except (ValueError, TypeError):
                                ws.write(
                                    row_idx, col_idx, str(val or ""), cell_fmt
                                )

                # Автофильтр
                if headers and rows:
                    ws.autofilter(0, 0, len(rows), len(headers) - 1)

            wb.close()
            return {"success": True, "path": filepath, "format": "xlsx"}

        except Exception as e:
            return {"error": f"Excel creation failed: {e}"}

    async def read(self, filepath: str) -> dict:
        """Прочитать Excel-файл."""
        try:
            import openpyxl

            wb = openpyxl.load_workbook(filepath, data_only=True)
            sheets = {}

            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                headers = [cell.value for cell in ws[1]] if ws.max_row else []
                rows = []

                for row in ws.iter_rows(min_row=2, values_only=True):
                    rows.append([v for v in row])

                sheets[sheet_name] = {
                    "headers": headers,
                    "rows": rows,
                    "row_count": len(rows),
                }

            wb.close()

            # Если один лист → упростить
            if len(sheets) == 1:
                data = list(sheets.values())[0]
                data["format"] = "xlsx"
                return data

            return {"sheets": sheets, "format": "xlsx"}

        except Exception as e:
            return {"error": f"Excel read failed: {e}"}


# ═══════════════════════════════════════════════════════════════════════════
# PDF Engine
# ═══════════════════════════════════════════════════════════════════════════

class PDFEngine:
    """Движок генерации PDF (инвойсы, отчёты)."""

    async def create(self, filepath: str, structure: dict) -> dict:
        """Создать PDF из структуры."""
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import (
                Paragraph,
                SimpleDocTemplate,
                Spacer,
                Table,
                TableStyle,
            )

            doc = SimpleDocTemplate(filepath, pagesize=A4)
            styles = getSampleStyleSheet()
            elements = []

            # Заголовок
            title = structure.get("title", "Документ")
            elements.append(Paragraph(title, styles["Title"]))
            elements.append(Spacer(1, 12))

            # Текстовый контент
            content = structure.get("content", "")
            if content:
                for line in content.split("\n"):
                    if line.strip():
                        elements.append(Paragraph(line, styles["Normal"]))
                        elements.append(Spacer(1, 6))

            # Таблица
            headers = structure.get("headers", [])
            rows = structure.get("rows", [])

            if headers:
                table_data = [headers]
                for row in rows:
                    table_data.append([str(v or "") for v in row])

                t = Table(table_data)
                t.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472C4")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTSIZE", (0, 0), (-1, 0), 10),
                    ("FONTSIZE", (0, 1), (-1, -1), 9),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("ALIGN", (0, 0), (-1, -1), "LEFT"),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1),
                     [colors.white, colors.HexColor("#D9E2F3")]),
                ]))
                elements.append(t)

            doc.build(elements)
            return {"success": True, "path": filepath, "format": "pdf"}

        except Exception as e:
            return {"error": f"PDF creation failed: {e}"}

    async def create_invoice(
        self,
        filepath: str,
        order_data: dict,
    ) -> dict:
        """
        Создать PDF-инвойс для заказа.
        По ТЗ: генератор документов → PDF-инвойсы по команде.
        """
        items = order_data.get("items", [])
        headers = ["#", "Наименование", "Кол-во", "Ед.", "Цена", "Сумма"]

        rows = []
        total = 0.0
        for i, item in enumerate(items, 1):
            qty = item.get("quantity", 0)
            price = item.get("unit_price", 0) or 0
            subtotal = qty * price
            total += subtotal

            rows.append([
                str(i), item.get("name", ""),
                str(qty), item.get("unit", "шт"),
                f"${price:.2f}", f"${subtotal:.2f}",
            ])

        rows.append(["", "", "", "", "ИТОГО:", f"${total:.2f}"])

        structure = {
            "title": f"Инвойс — Заказ #{order_data.get('order_number', '?')}",
            "headers": headers,
            "rows": rows,
            "content": (
                f"Дата: {datetime.now().strftime('%Y-%m-%d')}\n"
                f"Поставщик: {order_data.get('supplier', '—')}\n"
                f"Клиент: {order_data.get('client', '—')}\n"
            ),
        }

        return await self.create(filepath, structure)
