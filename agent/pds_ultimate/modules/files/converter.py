"""
PDS-Ultimate File Converter (Standalone)
===========================================
Конвертация между форматами файлов.

По ТЗ §5.2:
- Word ↔ PDF
- Excel ↔ CSV
- Конвертация между любыми поддерживаемыми форматами
- Пакетная конвертация

Поддерживаемые конвертации:
    xlsx → csv, pdf, txt
    csv  → xlsx, txt
    docx → pdf, txt
    pdf  → txt (extraction)
    txt  → pdf, docx
    json → csv, xlsx, txt
"""

from __future__ import annotations

import csv
import json
import os
from pathlib import Path
from typing import Optional

from pds_ultimate.config import USER_FILES_DIR

# ─── Conversion Matrix ──────────────────────────────────────────────────────

SUPPORTED_CONVERSIONS = {
    "xlsx": ["csv", "pdf", "txt", "json"],
    "csv": ["xlsx", "txt", "json"],
    "docx": ["pdf", "txt"],
    "pdf": ["txt"],
    "txt": ["pdf", "docx", "csv"],
    "json": ["csv", "xlsx", "txt"],
    "md": ["pdf", "txt", "docx"],
}


class ConversionResult:
    """Результат конвертации."""

    def __init__(
        self,
        success: bool = True,
        source_path: str = "",
        target_path: str = "",
        source_format: str = "",
        target_format: str = "",
        error: Optional[str] = None,
    ):
        self.success = success
        self.source_path = source_path
        self.target_path = target_path
        self.source_format = source_format
        self.target_format = target_format
        self.error = error

    def to_dict(self) -> dict:
        return {
            "success": self.success,
            "source": self.source_path,
            "target": self.target_path,
            "from": self.source_format,
            "to": self.target_format,
            "error": self.error,
        }


class FileConverter:
    """
    Конвертер файлов между форматами.

    Использование:
        result = await converter.convert("/path/to/file.xlsx", "csv")
        formats = converter.get_supported_formats("xlsx")
        can = converter.can_convert("xlsx", "csv")
    """

    def __init__(self):
        self._conversion_count = 0

    # ═══════════════════════════════════════════════════════════════════════
    # Public API
    # ═══════════════════════════════════════════════════════════════════════

    def can_convert(self, from_format: str, to_format: str) -> bool:
        """Проверить, поддерживается ли конвертация."""
        from_fmt = from_format.lower().lstrip(".")
        to_fmt = to_format.lower().lstrip(".")
        return to_fmt in SUPPORTED_CONVERSIONS.get(from_fmt, [])

    def get_supported_formats(self, from_format: str) -> list[str]:
        """Получить список форматов, в которые можно конвертировать."""
        return SUPPORTED_CONVERSIONS.get(
            from_format.lower().lstrip("."), []
        )

    async def convert(
        self,
        source_path: str,
        target_format: str,
        output_dir: Optional[str] = None,
    ) -> ConversionResult:
        """
        Конвертировать файл в другой формат.

        Args:
            source_path: Путь к исходному файлу
            target_format: Целевой формат (csv, pdf, xlsx, etc.)
            output_dir: Директория для результата (default: USER_FILES_DIR)
        """
        source = Path(source_path)
        if not source.exists():
            return ConversionResult(
                success=False,
                source_path=source_path,
                source_format=source.suffix.lstrip("."),
                target_format=target_format,
                error="Файл не найден",
            )

        source_format = source.suffix.lower().lstrip(".")
        target_format = target_format.lower().lstrip(".")

        if not self.can_convert(source_format, target_format):
            return ConversionResult(
                success=False,
                source_path=source_path,
                source_format=source_format,
                target_format=target_format,
                error=(
                    f"Конвертация {source_format} → {target_format} "
                    f"не поддерживается"
                ),
            )

        # Целевой путь
        out_dir = output_dir or str(USER_FILES_DIR)
        os.makedirs(out_dir, exist_ok=True)
        target_name = f"{source.stem}.{target_format}"
        target_path = os.path.join(out_dir, target_name)

        # Маршрутизация конвертации
        converter_key = f"{source_format}_to_{target_format}"
        converter_method = getattr(
            self, f"_convert_{converter_key}", None
        )

        if converter_method:
            return await converter_method(source_path, target_path)

        # Fallback через промежуточный текст
        return await self._convert_via_text(
            source_path, target_path, source_format, target_format
        )

    async def convert_batch(
        self,
        source_paths: list[str],
        target_format: str,
        output_dir: Optional[str] = None,
    ) -> list[ConversionResult]:
        """Пакетная конвертация."""
        results = []
        for path in source_paths:
            result = await self.convert(path, target_format, output_dir)
            results.append(result)
        return results

    # ═══════════════════════════════════════════════════════════════════════
    # Specific Converters
    # ═══════════════════════════════════════════════════════════════════════

    async def _convert_xlsx_to_csv(
        self, source: str, target: str
    ) -> ConversionResult:
        """Excel → CSV."""
        try:
            import openpyxl

            wb = openpyxl.load_workbook(source, data_only=True)
            ws = wb.active

            with open(target, "w", encoding="utf-8-sig", newline="") as f:
                writer = csv.writer(f)
                for row in ws.iter_rows(values_only=True):
                    writer.writerow(
                        [str(v) if v is not None else "" for v in row]
                    )

            wb.close()
            self._conversion_count += 1

            return ConversionResult(
                source_path=source,
                target_path=target,
                source_format="xlsx",
                target_format="csv",
            )

        except Exception as e:
            return ConversionResult(
                success=False,
                source_path=source,
                target_path=target,
                source_format="xlsx",
                target_format="csv",
                error=str(e),
            )

    async def _convert_csv_to_xlsx(
        self, source: str, target: str,
    ) -> ConversionResult:
        """CSV → Excel."""
        try:
            from pds_ultimate.modules.files.excel_engine import excel_engine

            with open(source, "r", encoding="utf-8-sig") as f:
                reader = csv.reader(f)
                rows = list(reader)

            headers = rows[0] if rows else []
            data = rows[1:] if len(rows) > 1 else []

            structure = {
                "title": Path(source).stem,
                "headers": headers,
                "rows": data,
            }

            result = await excel_engine.create(target, structure)
            self._conversion_count += 1

            if result.get("success"):
                return ConversionResult(
                    source_path=source,
                    target_path=target,
                    source_format="csv",
                    target_format="xlsx",
                )

            return ConversionResult(
                success=False,
                source_path=source,
                target_path=target,
                source_format="csv",
                target_format="xlsx",
                error=result.get("error", "Unknown"),
            )

        except Exception as e:
            return ConversionResult(
                success=False, source_path=source,
                target_path=target,
                source_format="csv", target_format="xlsx",
                error=str(e),
            )

    async def _convert_xlsx_to_pdf(
        self, source: str, target: str,
    ) -> ConversionResult:
        """Excel → PDF (таблица)."""
        try:
            from pds_ultimate.modules.files.excel_engine import excel_engine
            from pds_ultimate.modules.files.pdf_engine import pdf_engine

            data = await excel_engine.read(source)
            if "error" in data:
                return ConversionResult(
                    success=False, source_path=source,
                    target_path=target,
                    source_format="xlsx", target_format="pdf",
                    error=data["error"],
                )

            structure = {
                "title": Path(source).stem,
                "headers": data.get("headers", []),
                "rows": data.get("rows", []),
            }

            result = await pdf_engine.create(target, structure)
            self._conversion_count += 1

            if result.get("success"):
                return ConversionResult(
                    source_path=source, target_path=target,
                    source_format="xlsx", target_format="pdf",
                )

            return ConversionResult(
                success=False, source_path=source,
                target_path=target,
                source_format="xlsx", target_format="pdf",
                error=result.get("error", "Unknown"),
            )

        except Exception as e:
            return ConversionResult(
                success=False, source_path=source,
                target_path=target,
                source_format="xlsx", target_format="pdf",
                error=str(e),
            )

    async def _convert_docx_to_pdf(
        self, source: str, target: str,
    ) -> ConversionResult:
        """Word → PDF."""
        try:
            from docx import Document

            from pds_ultimate.modules.files.pdf_engine import pdf_engine

            doc = Document(source)
            paragraphs = [
                p.text for p in doc.paragraphs if p.text.strip()
            ]
            content = "\n".join(paragraphs)

            structure = {
                "title": Path(source).stem,
                "content": content,
            }

            result = await pdf_engine.create(target, structure)
            self._conversion_count += 1

            if result.get("success"):
                return ConversionResult(
                    source_path=source, target_path=target,
                    source_format="docx", target_format="pdf",
                )

            return ConversionResult(
                success=False, source_path=source,
                target_path=target,
                source_format="docx", target_format="pdf",
                error=result.get("error", "Unknown"),
            )

        except Exception as e:
            return ConversionResult(
                success=False, source_path=source,
                target_path=target,
                source_format="docx", target_format="pdf",
                error=str(e),
            )

    async def _convert_docx_to_txt(
        self, source: str, target: str,
    ) -> ConversionResult:
        """Word → Text."""
        try:
            from docx import Document

            doc = Document(source)
            text = "\n".join(
                p.text for p in doc.paragraphs if p.text.strip()
            )

            with open(target, "w", encoding="utf-8") as f:
                f.write(text)

            self._conversion_count += 1
            return ConversionResult(
                source_path=source, target_path=target,
                source_format="docx", target_format="txt",
            )

        except Exception as e:
            return ConversionResult(
                success=False, source_path=source,
                target_path=target,
                source_format="docx", target_format="txt",
                error=str(e),
            )

    async def _convert_pdf_to_txt(
        self, source: str, target: str,
    ) -> ConversionResult:
        """PDF → Text."""
        try:
            from pds_ultimate.modules.files.pdf_engine import pdf_engine

            data = await pdf_engine.read(source)
            if "error" in data:
                return ConversionResult(
                    success=False, source_path=source,
                    target_path=target,
                    source_format="pdf", target_format="txt",
                    error=data["error"],
                )

            with open(target, "w", encoding="utf-8") as f:
                f.write(data.get("content", ""))

            self._conversion_count += 1
            return ConversionResult(
                source_path=source, target_path=target,
                source_format="pdf", target_format="txt",
            )

        except Exception as e:
            return ConversionResult(
                success=False, source_path=source,
                target_path=target,
                source_format="pdf", target_format="txt",
                error=str(e),
            )

    async def _convert_json_to_csv(
        self, source: str, target: str,
    ) -> ConversionResult:
        """JSON → CSV."""
        try:
            with open(source, "r", encoding="utf-8") as f:
                data = json.load(f)

            if isinstance(data, list) and data:
                # List of dicts
                if isinstance(data[0], dict):
                    headers = list(data[0].keys())
                    rows = [
                        [str(item.get(h, "")) for h in headers]
                        for item in data
                    ]
                else:
                    headers = [f"col_{i}" for i in range(len(data[0]))]
                    rows = [[str(v) for v in row] for row in data]

                with open(
                    target, "w", encoding="utf-8-sig", newline=""
                ) as f:
                    writer = csv.writer(f)
                    writer.writerow(headers)
                    writer.writerows(rows)

                self._conversion_count += 1
                return ConversionResult(
                    source_path=source, target_path=target,
                    source_format="json", target_format="csv",
                )

            return ConversionResult(
                success=False, source_path=source,
                target_path=target,
                source_format="json", target_format="csv",
                error="JSON не является массивом объектов",
            )

        except Exception as e:
            return ConversionResult(
                success=False, source_path=source,
                target_path=target,
                source_format="json", target_format="csv",
                error=str(e),
            )

    async def _convert_txt_to_pdf(
        self, source: str, target: str,
    ) -> ConversionResult:
        """Text → PDF."""
        try:
            from pds_ultimate.modules.files.pdf_engine import pdf_engine

            with open(source, "r", encoding="utf-8") as f:
                content = f.read()

            structure = {
                "title": Path(source).stem,
                "content": content,
            }

            result = await pdf_engine.create(target, structure)
            self._conversion_count += 1

            if result.get("success"):
                return ConversionResult(
                    source_path=source, target_path=target,
                    source_format="txt", target_format="pdf",
                )

            return ConversionResult(
                success=False, source_path=source,
                target_path=target,
                source_format="txt", target_format="pdf",
                error=result.get("error"),
            )

        except Exception as e:
            return ConversionResult(
                success=False, source_path=source,
                target_path=target,
                source_format="txt", target_format="pdf",
                error=str(e),
            )

    # ═══════════════════════════════════════════════════════════════════════
    # Fallback: via text
    # ═══════════════════════════════════════════════════════════════════════

    async def _convert_via_text(
        self,
        source: str,
        target: str,
        source_format: str,
        target_format: str,
    ) -> ConversionResult:
        """Fallback: конвертация через промежуточный текст."""
        try:
            # Читаем как текст
            with open(source, "r", encoding="utf-8") as f:
                content = f.read()

            # Пишем в целевой формат
            if target_format == "txt":
                with open(target, "w", encoding="utf-8") as f:
                    f.write(content)
            elif target_format == "json":
                with open(target, "w", encoding="utf-8") as f:
                    json.dump(
                        {"content": content},
                        f, ensure_ascii=False, indent=2,
                    )
            else:
                return ConversionResult(
                    success=False, source_path=source,
                    target_path=target,
                    source_format=source_format,
                    target_format=target_format,
                    error="Нет подходящего конвертера",
                )

            self._conversion_count += 1
            return ConversionResult(
                source_path=source, target_path=target,
                source_format=source_format,
                target_format=target_format,
            )

        except Exception as e:
            return ConversionResult(
                success=False, source_path=source,
                target_path=target,
                source_format=source_format,
                target_format=target_format,
                error=str(e),
            )

    # ═══════════════════════════════════════════════════════════════════════
    # Formatting
    # ═══════════════════════════════════════════════════════════════════════

    def format_result(self, result: ConversionResult) -> str:
        """Форматировать результат для бота."""
        if result.success:
            return (
                f"✅ Конвертация {result.source_format.upper()} → "
                f"{result.target_format.upper()}\n"
                f"📄 {Path(result.target_path).name}"
            )
        return (
            f"❌ Ошибка конвертации: {result.error}"
        )

    def format_supported(self, from_format: str) -> str:
        """Форматировать список поддерживаемых конвертаций."""
        formats = self.get_supported_formats(from_format)
        if formats:
            return (
                f"📎 {from_format.upper()} можно конвертировать в: "
                + ", ".join(f.upper() for f in formats)
            )
        return f"❌ Формат {from_format} не поддерживается"

    def get_stats(self) -> dict:
        return {
            "total_conversions": self._conversion_count,
            "supported_formats": list(SUPPORTED_CONVERSIONS.keys()),
        }


# ─── Глобальный экземпляр ────────────────────────────────────────────────────

file_converter = FileConverter()
