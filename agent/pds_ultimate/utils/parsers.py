"""
PDS-Ultimate Universal Parser
================================
Универсальный парсер данных.

По ТЗ принимает заказы в ЛЮБОМ виде:
- Текст: «Балаклавы 100 шт, цена 200$»
- Excel (.xlsx, .xls)
- Word (.docx)
- PDF
- Фото накладной/чека (OCR)
- Голосовое сообщение (Faster-Whisper → текст → парсинг)
- CSV

DeepSeek + regex одновременно для 100% точности.
Мгновенная склейка: файл + текстовое дополнение → один заказ.
"""

from __future__ import annotations

import os
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from pds_ultimate.config import config, logger

# ─── Структуры данных ────────────────────────────────────────────────────────


@dataclass
class ParsedItem:
    """Распознанная позиция товара."""
    name: str
    quantity: float
    unit: str = "шт"
    unit_price: Optional[float] = None
    currency: str = "USD"
    weight: Optional[float] = None
    weight_unit: str = "кг"
    notes: Optional[str] = None

    def to_dict(self) -> dict:
        return {
            "name": self.name,
            "quantity": self.quantity,
            "unit": self.unit,
            "unit_price": self.unit_price,
            "currency": self.currency,
            "weight": self.weight,
            "weight_unit": self.weight_unit,
            "notes": self.notes,
        }


@dataclass
class ParseResult:
    """Результат парсинга."""
    items: list[ParsedItem] = field(default_factory=list)
    raw_text: str = ""
    source_type: str = ""  # text, excel, word, pdf, image, voice
    source_file: Optional[str] = None
    errors: list[str] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)

    @property
    def success(self) -> bool:
        return len(self.items) > 0 and len(self.errors) == 0

    @property
    def item_count(self) -> int:
        return len(self.items)

    def merge(self, other: "ParseResult") -> "ParseResult":
        """Объединить два результата (для склейки файл + текст)."""
        merged = ParseResult(
            items=self.items + other.items,
            raw_text=f"{self.raw_text}\n---\n{other.raw_text}",
            source_type=f"{self.source_type}+{other.source_type}",
            errors=self.errors + other.errors,
            warnings=self.warnings + other.warnings,
        )
        return merged


# ─── Regex-парсер ────────────────────────────────────────────────────────────

class RegexParser:
    """
    Парсинг текста через регулярные выражения.
    Первый слой обработки — быстрый и точный для простых форматов.
    """

    # Паттерны для распознавания товаров
    # Форматы: "Балаклавы 100 шт", "100 шт балаклав", "Маски - 50 шт по $2.5"
    PATTERNS = [
        # "Балаклавы 100 шт по 200$"
        re.compile(
            r"(?P<name>[А-Яа-яёЁA-Za-z\s\-]+?)\s+"
            r"(?P<qty>\d+[.,]?\d*)\s*"
            r"(?P<unit>шт|кг|м|л|уп|кор|пач|рул|компл)\.?"
            r"(?:\s*(?:по|@|x|х)\s*"
            r"(?P<price>\d+[.,]?\d*)\s*"
            r"(?P<currency>\$|USD|¥|CNY|юан|ман|TMT|€|EUR)?)?",
            re.IGNORECASE | re.UNICODE,
        ),
        # "100 шт Балаклав по $200"
        re.compile(
            r"(?P<qty>\d+[.,]?\d*)\s*"
            r"(?P<unit>шт|кг|м|л|уп|кор|пач|рул|компл)\.?\s+"
            r"(?P<name>[А-Яа-яёЁA-Za-z\s\-]+?)"
            r"(?:\s*(?:по|@|x|х)\s*"
            r"(?P<price>\d+[.,]?\d*)\s*"
            r"(?P<currency>\$|USD|¥|CNY|юан|ман|TMT|€|EUR)?)?",
            re.IGNORECASE | re.UNICODE,
        ),
    ]

    # Маппинг символов валют
    CURRENCY_MAP = {
        "$": "USD",
        "usd": "USD",
        "¥": "CNY",
        "cny": "CNY",
        "юан": "CNY",
        "юань": "CNY",
        "ман": "TMT",
        "манат": "TMT",
        "tmt": "TMT",
        "€": "EUR",
        "eur": "EUR",
        "руб": "RUB",
        "rub": "RUB",
        "₽": "RUB",
    }

    UNIT_MAP = {
        "шт": "шт",
        "штук": "шт",
        "штуки": "шт",
        "кг": "кг",
        "килограмм": "кг",
        "м": "м",
        "метр": "м",
        "л": "л",
        "литр": "л",
        "уп": "уп",
        "упаковка": "уп",
        "кор": "кор",
        "коробка": "кор",
        "пач": "пач",
        "пачка": "пач",
        "рул": "рул",
        "рулон": "рул",
        "компл": "компл",
        "комплект": "компл",
    }

    @classmethod
    def parse(cls, text: str) -> list[ParsedItem]:
        """Парсинг текста через regex."""
        items = []
        lines = text.strip().split("\n")

        for line in lines:
            line = line.strip()
            if not line or line.startswith("#") or line.startswith("//"):
                continue

            for pattern in cls.PATTERNS:
                match = pattern.search(line)
                if match:
                    groups = match.groupdict()

                    name = groups.get("name", "").strip().rstrip(" -,.")
                    if not name or len(name) < 2:
                        continue

                    qty_str = groups.get("qty", "0").replace(",", ".")
                    try:
                        qty = float(qty_str)
                    except ValueError:
                        continue

                    if qty <= 0:
                        continue

                    unit = cls.UNIT_MAP.get(
                        groups.get("unit", "шт").lower().rstrip("."),
                        "шт"
                    )

                    price = None
                    price_str = groups.get("price")
                    if price_str:
                        try:
                            price = float(price_str.replace(",", "."))
                        except ValueError:
                            pass

                    currency = "USD"
                    curr_str = groups.get("currency")
                    if curr_str:
                        currency = cls.CURRENCY_MAP.get(
                            curr_str.lower().rstrip("."),
                            "USD"
                        )

                    items.append(ParsedItem(
                        name=name.capitalize(),
                        quantity=qty,
                        unit=unit,
                        unit_price=price,
                        currency=currency,
                    ))
                    break  # Один матч на строку

        return items


# ─── Excel-парсер ────────────────────────────────────────────────────────────

class ExcelParser:
    """Парсинг Excel-файлов."""

    @staticmethod
    def parse(file_path: str | Path) -> ParseResult:
        """Парсинг Excel-файла."""
        try:
            import pandas as pd

            file_path = Path(file_path)
            result = ParseResult(source_type="excel",
                                 source_file=str(file_path))

            # Читаем все листы
            xl = pd.ExcelFile(file_path)
            all_items = []

            for sheet_name in xl.sheet_names:
                df = pd.read_excel(xl, sheet_name=sheet_name)

                if df.empty:
                    continue

                # Нормализация колонок
                col_map = ExcelParser._detect_columns(df)

                if not col_map.get("name"):
                    result.warnings.append(
                        f"Лист '{sheet_name}': не найдена колонка с названием товара"
                    )
                    # Попробовать весь лист как текст
                    text = df.to_string()
                    regex_items = RegexParser.parse(text)
                    all_items.extend(regex_items)
                    continue

                for _, row in df.iterrows():
                    name = str(row.get(col_map["name"], "")).strip()
                    if not name or name.lower() in ("nan", "", "none"):
                        continue

                    qty = 1.0
                    qty_col = col_map.get("quantity")
                    if qty_col and pd.notna(row.get(qty_col)):
                        try:
                            qty = float(row[qty_col])
                        except (ValueError, TypeError):
                            pass

                    price = None
                    price_col = col_map.get("price")
                    if price_col and pd.notna(row.get(price_col)):
                        try:
                            price = float(row[price_col])
                        except (ValueError, TypeError):
                            pass

                    unit = "шт"
                    unit_col = col_map.get("unit")
                    if unit_col and pd.notna(row.get(unit_col)):
                        unit_raw = str(row[unit_col]).lower().strip()
                        unit = RegexParser.UNIT_MAP.get(unit_raw, unit_raw)

                    weight = None
                    weight_col = col_map.get("weight")
                    if weight_col and pd.notna(row.get(weight_col)):
                        try:
                            weight = float(row[weight_col])
                        except (ValueError, TypeError):
                            pass

                    all_items.append(ParsedItem(
                        name=name,
                        quantity=qty,
                        unit=unit,
                        unit_price=price,
                        weight=weight,
                    ))

            result.items = all_items
            result.raw_text = f"[Excel: {file_path.name}, {len(all_items)} позиций]"
            return result

        except ImportError:
            return ParseResult(
                source_type="excel",
                errors=["Библиотека pandas/openpyxl не установлена"],
            )
        except Exception as e:
            logger.error(f"Ошибка парсинга Excel: {e}")
            return ParseResult(
                source_type="excel",
                errors=[f"Ошибка чтения Excel: {str(e)}"],
            )

    @staticmethod
    def _detect_columns(df) -> dict:
        """Автоопределение колонок по названиям."""
        col_map = {}
        columns_lower = {col: col.lower().strip() for col in df.columns}

        name_variants = [
            "название", "наименование", "товар", "name", "item",
            "product", "позиция", "описание", "description",
        ]
        qty_variants = [
            "количество", "кол-во", "кол", "qty", "quantity",
            "count", "шт", "число",
        ]
        price_variants = [
            "цена", "price", "стоимость", "cost", "сумма",
            "unit_price", "unit price",
        ]
        unit_variants = [
            "единица", "ед", "unit", "ед.изм", "ед. изм.",
            "единица измерения",
        ]
        weight_variants = [
            "вес", "weight", "масса", "кг", "kg",
        ]

        for col, col_lower in columns_lower.items():
            for variant in name_variants:
                if variant in col_lower:
                    col_map["name"] = col
                    break
            for variant in qty_variants:
                if variant in col_lower:
                    col_map["quantity"] = col
                    break
            for variant in price_variants:
                if variant in col_lower:
                    col_map["price"] = col
                    break
            for variant in unit_variants:
                if variant in col_lower:
                    col_map["unit"] = col
                    break
            for variant in weight_variants:
                if variant in col_lower:
                    col_map["weight"] = col
                    break

        # Если не нашли по имени — берём первую текстовую колонку как "name"
        if "name" not in col_map:
            for col in df.columns:
                if df[col].dtype == object:
                    col_map["name"] = col
                    break

        return col_map


# ─── Word-парсер ─────────────────────────────────────────────────────────────

class WordParser:
    """Парсинг Word-документов (.docx)."""

    @staticmethod
    def parse(file_path: str | Path) -> ParseResult:
        """Парсинг Word-файла — таблицы + текст."""
        try:
            from docx import Document

            file_path = Path(file_path)
            result = ParseResult(source_type="word",
                                 source_file=str(file_path))
            doc = Document(file_path)
            all_text_parts = []
            all_items = []

            # 1. Парсим таблицы
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    row_text = " | ".join(cells)
                    all_text_parts.append(row_text)

            # 2. Парсим параграфы
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    all_text_parts.append(text)

            full_text = "\n".join(all_text_parts)
            result.raw_text = full_text

            # Парсим через regex
            all_items = RegexParser.parse(full_text)
            result.items = all_items

            return result

        except ImportError:
            return ParseResult(
                source_type="word",
                errors=["Библиотека python-docx не установлена"],
            )
        except Exception as e:
            logger.error(f"Ошибка парсинга Word: {e}")
            return ParseResult(
                source_type="word",
                errors=[f"Ошибка чтения Word: {str(e)}"],
            )


# ─── PDF-парсер ──────────────────────────────────────────────────────────────

class PDFParser:
    """Парсинг PDF-документов."""

    @staticmethod
    def parse(file_path: str | Path) -> ParseResult:
        """Парсинг PDF-файла."""
        try:
            from PyPDF2 import PdfReader

            file_path = Path(file_path)
            result = ParseResult(source_type="pdf", source_file=str(file_path))

            reader = PdfReader(file_path)
            all_text_parts = []

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    all_text_parts.append(text)

            full_text = "\n".join(all_text_parts)
            result.raw_text = full_text

            if not full_text.strip():
                result.warnings.append(
                    "PDF не содержит текстового слоя. Попробуйте OCR."
                )
                return result

            result.items = RegexParser.parse(full_text)
            return result

        except ImportError:
            return ParseResult(
                source_type="pdf",
                errors=["Библиотека PyPDF2 не установлена"],
            )
        except Exception as e:
            logger.error(f"Ошибка парсинга PDF: {e}")
            return ParseResult(
                source_type="pdf",
                errors=[f"Ошибка чтения PDF: {str(e)}"],
            )


# ─── OCR-парсер (Фото) ──────────────────────────────────────────────────────

class OCRParser:
    """
    Парсинг изображений через OCR.
    Для: фото накладных, чеков, трек-номеров.
    """

    @staticmethod
    def parse(file_path: str | Path) -> ParseResult:
        """Парсинг изображения через OCR."""
        result = ParseResult(source_type="image", source_file=str(file_path))

        text = OCRParser.extract_text(str(file_path))
        if not text:
            result.errors.append("Не удалось распознать текст на изображении")
            return result

        result.raw_text = text
        result.items = RegexParser.parse(text)
        return result

    @staticmethod
    def extract_text(file_path: str) -> str:
        """
        Извлечь текст из изображения.
        Используется также для извлечения трек-номеров.
        """
        engine = config.ocr.engine

        if engine == "easyocr":
            return OCRParser._extract_easyocr(file_path)
        elif engine == "tesseract":
            return OCRParser._extract_tesseract(file_path)
        else:
            logger.error(f"Неизвестный OCR движок: {engine}")
            return ""

    @staticmethod
    def extract_tracking_number(file_path: str) -> Optional[str]:
        """
        Извлечь трек-номер из фото.
        По ТЗ: если скинули фото чека с треком — номер извлекается автоматически.
        """
        text = OCRParser.extract_text(file_path)
        if not text:
            return None

        # Паттерны трек-номеров
        patterns = [
            # Китай: SF, YT, EMS, EY, etc.
            r"(?:SF|YT|EMS|EY|LP|JT)\d{10,15}",
            # DHL, UPS, FedEx
            r"\b\d{10,22}\b",
            # Общий: буквы + цифры, 10-20 символов
            r"\b[A-Z]{2}\d{9,13}[A-Z]{2}\b",
        ]

        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(0)

        return None

    @staticmethod
    def _extract_easyocr(file_path: str) -> str:
        """OCR через EasyOCR."""
        try:
            import easyocr

            reader = easyocr.Reader(
                config.ocr.languages,
                gpu=config.whisper.device != "cpu",
            )
            results = reader.readtext(file_path)

            texts = []
            for (bbox, text, confidence) in results:
                if confidence >= config.ocr.confidence_threshold:
                    texts.append(text)

            return "\n".join(texts)

        except ImportError:
            logger.error("EasyOCR не установлен: pip install easyocr")
            return ""
        except Exception as e:
            logger.error(f"Ошибка EasyOCR: {e}")
            return ""

    @staticmethod
    def _extract_tesseract(file_path: str) -> str:
        """OCR через Tesseract."""
        try:
            import pytesseract
            from PIL import Image

            image = Image.open(file_path)
            lang = "+".join(config.ocr.languages)
            text = pytesseract.image_to_string(image, lang=lang)
            return text

        except ImportError:
            logger.error(
                "Tesseract не установлен: pip install pytesseract Pillow"
            )
            return ""
        except Exception as e:
            logger.error(f"Ошибка Tesseract: {e}")
            return ""


# ─── Voice-парсер (Faster-Whisper) ──────────────────────────────────────────

class VoiceParser:
    """
    Распознавание голосовых сообщений через Faster-Whisper (локально).
    По ТЗ: бесплатно, работает на GPU/CPU сервера.
    """

    _model = None  # Ленивая загрузка модели

    @classmethod
    def _load_model(cls):
        """Загрузить модель Whisper (один раз)."""
        if cls._model is not None:
            return

        try:
            from faster_whisper import WhisperModel

            device = config.whisper.device
            if device == "auto":
                try:
                    import torch
                    device = "cuda" if torch.cuda.is_available() else "cpu"
                except ImportError:
                    device = "cpu"

            model_dir = str(config.whisper.model_dir)
            os.makedirs(model_dir, exist_ok=True)

            cls._model = WhisperModel(
                config.whisper.model_size,
                device=device,
                compute_type=config.whisper.compute_type,
                download_root=model_dir,
            )
            logger.info(
                f"Faster-Whisper загружен: model={config.whisper.model_size}, "
                f"device={device}"
            )
        except ImportError:
            logger.error(
                "Faster-Whisper не установлен: pip install faster-whisper"
            )
            raise
        except Exception as e:
            logger.error(f"Ошибка загрузки Whisper: {e}")
            raise

    @classmethod
    def transcribe(cls, audio_path: str) -> str:
        """
        Транскрибировать аудио в текст.
        Поддерживает: ogg, mp3, wav, m4a и другие форматы.
        """
        cls._load_model()

        try:
            segments, info = cls._model.transcribe(
                audio_path,
                language=config.whisper.language,
                beam_size=5,
                vad_filter=True,  # Фильтрация тишины
                vad_parameters=dict(
                    min_silence_duration_ms=500,
                ),
            )

            text_parts = []
            for segment in segments:
                text_parts.append(segment.text.strip())

            full_text = " ".join(text_parts)
            logger.info(
                f"Whisper: распознано {len(text_parts)} сегментов, "
                f"язык={info.language}, вероятность={info.language_probability:.2f}"
            )
            return full_text

        except Exception as e:
            logger.error(f"Ошибка транскрибации: {e}")
            return ""

    @classmethod
    def parse(cls, audio_path: str) -> ParseResult:
        """Распознать голос и парсить как текст."""
        result = ParseResult(source_type="voice", source_file=audio_path)

        text = cls.transcribe(audio_path)
        if not text:
            result.errors.append("Не удалось распознать голосовое сообщение")
            return result

        result.raw_text = text
        result.items = RegexParser.parse(text)
        return result


# ─── CSV-парсер ──────────────────────────────────────────────────────────────

class CSVParser:
    """Парсинг CSV-файлов."""

    @staticmethod
    def parse(file_path: str | Path) -> ParseResult:
        """Парсинг CSV-файла (делегирует в Excel-парсер через pandas)."""
        try:
            import pandas as pd

            file_path = Path(file_path)
            result = ParseResult(source_type="csv", source_file=str(file_path))

            # Пробуем разные разделители
            for sep in [",", ";", "\t", "|"]:
                try:
                    df = pd.read_csv(file_path, sep=sep)
                    if len(df.columns) > 1:
                        break
                except Exception:
                    continue

            # Используем логику Excel-парсера для определения колонок
            col_map = ExcelParser._detect_columns(df)

            if col_map.get("name"):
                for _, row in df.iterrows():
                    name = str(row.get(col_map["name"], "")).strip()
                    if not name or name.lower() in ("nan", ""):
                        continue

                    qty = 1.0
                    qty_col = col_map.get("quantity")
                    if qty_col and pd.notna(row.get(qty_col)):
                        try:
                            qty = float(row[qty_col])
                        except (ValueError, TypeError):
                            pass

                    price = None
                    price_col = col_map.get("price")
                    if price_col and pd.notna(row.get(price_col)):
                        try:
                            price = float(row[price_col])
                        except (ValueError, TypeError):
                            pass

                    result.items.append(ParsedItem(
                        name=name,
                        quantity=qty,
                        unit_price=price,
                    ))
            else:
                # Fallback: regex парсинг
                text = df.to_string()
                result.items = RegexParser.parse(text)

            result.raw_text = f"[CSV: {file_path.name}, {len(result.items)} позиций]"
            return result

        except Exception as e:
            logger.error(f"Ошибка парсинга CSV: {e}")
            return ParseResult(
                source_type="csv",
                errors=[f"Ошибка чтения CSV: {str(e)}"],
            )


# ─── Главный парсер (фасад) ─────────────────────────────────────────────────

class UniversalParser:
    """
    Универсальный парсер — единая точка входа для всех форматов.

    Использование:
        parser = UniversalParser()
        result = parser.parse_text("Балаклавы 100 шт, маски 50 шт")
        result = parser.parse_file("/path/to/order.xlsx")
        result = parser.parse_voice("/path/to/voice.ogg")
        result = parser.parse_image("/path/to/photo.jpg")
    """

    # Маппинг расширений на парсеры
    FILE_PARSERS = {
        ".xlsx": ExcelParser,
        ".xls": ExcelParser,
        ".csv": CSVParser,
        ".docx": WordParser,
        ".pdf": PDFParser,
        ".jpg": OCRParser,
        ".jpeg": OCRParser,
        ".png": OCRParser,
        ".bmp": OCRParser,
        ".tiff": OCRParser,
        ".webp": OCRParser,
    }

    VOICE_EXTENSIONS = {".ogg", ".mp3", ".wav", ".m4a", ".flac", ".opus"}

    def __init__(self):
        self._llm_engine = None  # Lazy loading

    async def _get_llm(self):
        """Получить LLM движок (lazy)."""
        if self._llm_engine is None:
            from pds_ultimate.core.llm_engine import llm_engine
            self._llm_engine = llm_engine
        return self._llm_engine

    # ─── Публичные методы ────────────────────────────────────────────────

    def parse_text(self, text: str) -> ParseResult:
        """Парсинг текстового ввода (regex)."""
        result = ParseResult(source_type="text", raw_text=text)
        result.items = RegexParser.parse(text)
        return result

    async def parse_text_smart(self, text: str) -> ParseResult:
        """
        Умный парсинг текста: regex + DeepSeek.
        Сначала regex, если мало результатов — подключаем LLM.
        """
        # Шаг 1: Regex
        result = self.parse_text(text)

        # Шаг 2: Если regex не справился — LLM
        if not result.items or len(text) > 200:
            llm = await self._get_llm()
            try:
                llm_items = await llm.parse_order(text)
                llm_parsed = [
                    ParsedItem(
                        name=item.get("name", ""),
                        quantity=float(item.get("quantity", 1)),
                        unit=item.get("unit", "шт"),
                        unit_price=item.get("unit_price"),
                        currency=item.get("currency", "USD"),
                    )
                    for item in llm_items
                    if item.get("name")
                ]

                if len(llm_parsed) > len(result.items):
                    result.items = llm_parsed
                    result.warnings.append("Использован LLM для парсинга")

            except Exception as e:
                logger.warning(f"LLM парсинг не удался: {e}")
                result.warnings.append(f"LLM парсинг не удался: {e}")

        return result

    def parse_file(self, file_path: str | Path) -> ParseResult:
        """
        Парсинг файла (автоопределение формата по расширению).
        """
        file_path = Path(file_path)

        if not file_path.exists():
            return ParseResult(
                source_type="file",
                errors=[f"Файл не найден: {file_path}"],
            )

        ext = file_path.suffix.lower()

        # Голосовое сообщение
        if ext in self.VOICE_EXTENSIONS:
            return VoiceParser.parse(str(file_path))

        # Файл
        parser_class = self.FILE_PARSERS.get(ext)
        if parser_class:
            return parser_class.parse(file_path)

        # Неизвестный формат — пробуем как текст
        try:
            text = file_path.read_text(encoding="utf-8")
            result = self.parse_text(text)
            result.source_type = "text_file"
            result.source_file = str(file_path)
            return result
        except Exception:
            return ParseResult(
                source_type="unknown",
                errors=[f"Неподдерживаемый формат файла: {ext}"],
            )

    async def parse_file_smart(self, file_path: str | Path) -> ParseResult:
        """
        Умный парсинг файла: формат-парсер + DeepSeek для проверки.
        """
        result = self.parse_file(file_path)

        # Если файл дал текст но мало позиций — подключаем LLM
        if result.raw_text and not result.items:
            llm = await self._get_llm()
            try:
                llm_items = await llm.parse_order(result.raw_text[:3000])
                result.items = [
                    ParsedItem(
                        name=item.get("name", ""),
                        quantity=float(item.get("quantity", 1)),
                        unit=item.get("unit", "шт"),
                        unit_price=item.get("unit_price"),
                        currency=item.get("currency", "USD"),
                    )
                    for item in llm_items
                    if item.get("name")
                ]
                result.warnings.append("Использован LLM для парсинга файла")
            except Exception as e:
                logger.warning(f"LLM парсинг файла не удался: {e}")

        return result

    def parse_voice(self, audio_path: str) -> ParseResult:
        """Парсинг голосового сообщения."""
        return VoiceParser.parse(audio_path)

    def parse_image(self, image_path: str) -> ParseResult:
        """Парсинг изображения (OCR)."""
        return OCRParser.parse(image_path)

    def extract_tracking_number(self, image_path: str) -> Optional[str]:
        """Извлечь трек-номер из фото."""
        return OCRParser.extract_tracking_number(image_path)

    def transcribe_voice(self, audio_path: str) -> str:
        """Транскрибировать голос в текст (без парсинга заказа)."""
        return VoiceParser.transcribe(audio_path)

    def detect_format(self, file_path: str | Path) -> str:
        """Определить формат файла."""
        ext = Path(file_path).suffix.lower()
        if ext in self.VOICE_EXTENSIONS:
            return "voice"
        if ext in self.FILE_PARSERS:
            return ext.lstrip(".")
        return "unknown"


# ─── Глобальный экземпляр ────────────────────────────────────────────────────

parser = UniversalParser()
